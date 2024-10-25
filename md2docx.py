# md文件转docx文件示例
from docx import Document
from markdown import markdown
from docx.shared import Inches
from bs4 import BeautifulSoup

def md_to_docx(md_file, docx_file):
    # 读取md文件
    try:
        with open(md_file, 'r', encoding='utf-8') as file:
            md_content = file.read()
        print(f"成功读取 {md_file}，内容长度：{len(md_content)} 字符")
    except Exception as e:
        print(f"读取 {md_file} 时出错：{e}")
        return

    # 将md内容转换为html
    html = markdown(md_content)
    print(f"Markdown 转换为 HTML，长度：{len(html)} 字符")

    # 创建一个新的Document对象
    doc = Document()

    # 使用BeautifulSoup解析html
    soup = BeautifulSoup(html, 'html.parser')

    # 遍历html元素并添加到docx
    element_count = 0
    for element in soup.find_all():
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        element_count += 1

    print(f"处理了 {element_count} 个 HTML 元素")

    # 保存docx文件
    try:
        doc.save(docx_file)
        print(f"成功保存 {docx_file}")
    except Exception as e:
        print(f"保存 {docx_file} 时出错：{e}")

# 使用示例
md_to_docx('example.md', 'output.docx')

print("转换完成")
